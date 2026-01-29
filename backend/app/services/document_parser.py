"""
Document parsing service for extracting text from various file formats
"""

import io
import logging
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile
import os

# PDF processing
try:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        # Fallback to PyPDF2 if pypdf not available
        from PyPDF2 import PdfReader
        from PyPDF2.errors import PdfReadError
        PYPDF_AVAILABLE = True
        logging.warning("Using deprecated PyPDF2 - consider upgrading to pypdf")
    except ImportError:
        PYPDF_AVAILABLE = False
        logging.warning("No PDF library available - PDF processing disabled")

# Alternative PDF processing with pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available - advanced PDF processing disabled")

# DOC/DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX processing disabled")

from fastapi import HTTPException, UploadFile

logger = logging.getLogger(__name__)

class DocumentParsingError(Exception):
    """Custom exception for document parsing errors"""
    pass

class DocumentParser:
    """Service for parsing various document formats and extracting text content"""
    
    def __init__(self):
        """Initialize document parser"""
        self.supported_formats = {}
        
        # Add PDF support if available
        if PYPDF_AVAILABLE or PDFPLUMBER_AVAILABLE:
            self.supported_formats['.pdf'] = self._extract_pdf_text
        
        # Add text file support
        self.supported_formats['.txt'] = self._extract_text_file
        
        # Add DOCX support if available
        if DOCX_AVAILABLE:
            self.supported_formats['.docx'] = self._extract_docx_text
            
        # DOC support (limited)
        self.supported_formats['.doc'] = self._extract_doc_text
        
        logger.info(f"Document parser initialized with support for: {list(self.supported_formats.keys())}")
    
    async def extract_text_from_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Extract text content from uploaded file
        
        Args:
            file: Uploaded file object
            
        Returns:
            Dictionary containing extracted text and metadata
            
        Raises:
            DocumentParsingError: If parsing fails
            HTTPException: If file format is unsupported
        """
        try:
            # Get file extension
            file_ext = Path(file.filename).suffix.lower()
            
            if file_ext not in self.supported_formats:
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file format: {file_ext}. Supported formats: {', '.join(self.supported_formats.keys())}"
                )
            
            # Read file content
            content = await file.read()
            
            # Reset file pointer for potential reuse
            await file.seek(0)
            
            # Validate file size (reasonable limits)
            max_file_size = 50 * 1024 * 1024  # 50MB
            if len(content) > max_file_size:
                raise HTTPException(
                    status_code=413,
                    detail=f"File too large: {len(content)} bytes. Maximum allowed: {max_file_size} bytes"
                )
            
            # Check for empty files
            if len(content) == 0:
                raise DocumentParsingError(f"File {file.filename} is empty")
            
            # Validate file format by checking magic bytes for common formats
            self._validate_file_format(content, file_ext, file.filename)
            
            # Extract text using appropriate parser
            parser_func = self.supported_formats[file_ext]
            extracted_text, metadata = await parser_func(content, file.filename)
            
            # Validate extracted content
            if not extracted_text or len(extracted_text.strip()) < 10:
                raise DocumentParsingError(
                    f"Insufficient text content extracted from {file.filename}. "
                    "Please ensure the file contains readable text."
                )
            
            result = {
                "text": extracted_text,
                "metadata": {
                    "filename": file.filename,
                    "file_size": len(content),
                    "file_type": file_ext,
                    "character_count": len(extracted_text),
                    "word_count": len(extracted_text.split()),
                    **metadata
                }
            }
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from {file.filename}")
            return result
            
        except HTTPException:
            raise
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing {file.filename}: {str(e)}")
            raise DocumentParsingError(f"Failed to parse document: {str(e)}")
    
    def _validate_file_format(self, content: bytes, expected_ext: str, filename: str):
        """
        Validate file format by checking magic bytes
        
        Args:
            content: File content as bytes
            expected_ext: Expected file extension
            filename: Original filename
            
        Raises:
            DocumentParsingError: If file format doesn't match extension
        """
        try:
            # Check magic bytes for common formats
            if expected_ext == '.pdf':
                if not content.startswith(b'%PDF'):
                    raise DocumentParsingError(
                        f"File {filename} does not appear to be a valid PDF file"
                    )
            
            elif expected_ext == '.docx':
                # DOCX files are ZIP archives, should start with PK
                if not content.startswith(b'PK'):
                    raise DocumentParsingError(
                        f"File {filename} does not appear to be a valid DOCX file"
                    )
            
            elif expected_ext == '.doc':
                # DOC files have specific signatures
                if not (content.startswith(b'\xd0\xcf\x11\xe0') or content.startswith(b'PK')):
                    # Allow PK for misnamed DOCX files
                    logger.warning(f"File {filename} may not be a valid DOC file")
            
            # For .txt files, we're more lenient as they can have various encodings
            
        except Exception as e:
            logger.warning(f"File format validation failed for {filename}: {str(e)}")
            # Don't raise error for validation failures, just log warning
    
    async def _extract_pdf_text(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file using multiple extraction methods
        
        Args:
            content: PDF file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            DocumentParsingError: If PDF parsing fails
        """
        if not PYPDF_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise DocumentParsingError("No PDF processing library available")
        
        # Try multiple extraction methods for best results
        extraction_methods = []
        
        if PDFPLUMBER_AVAILABLE:
            extraction_methods.append(("pdfplumber", self._extract_pdf_with_pdfplumber))
        
        if PYPDF_AVAILABLE:
            extraction_methods.append(("pypdf", self._extract_pdf_with_pypdf))
        
        last_error = None
        
        for method_name, method_func in extraction_methods:
            try:
                logger.info(f"Attempting PDF extraction with {method_name} for {filename}")
                extracted_text, metadata = await method_func(content, filename)
                
                # If we got good text, return it
                if extracted_text and len(extracted_text.strip()) >= 10:
                    metadata["extraction_method"] = method_name
                    logger.info(f"Successfully extracted {len(extracted_text)} characters using {method_name}")
                    return extracted_text, metadata
                else:
                    logger.warning(f"{method_name} extracted insufficient text from {filename}")
                    
            except Exception as e:
                logger.warning(f"{method_name} extraction failed for {filename}: {str(e)}")
                last_error = e
                continue
        
        # If all methods failed
        if last_error:
            raise DocumentParsingError(f"All PDF extraction methods failed for {filename}. Last error: {str(last_error)}")
        else:
            raise DocumentParsingError(f"No readable text found in PDF {filename} using any extraction method")
    
    async def _extract_pdf_with_pdfplumber(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using pdfplumber (better for complex layouts)
        
        Args:
            content: PDF file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Create a temporary file for pdfplumber
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                extracted_text = []
                metadata = {
                    "page_count": 0,
                    "successful_pages": 0,
                    "failed_pages": 0,
                    "tables_found": 0,
                    "extraction_method": "pdfplumber"
                }
                
                with pdfplumber.open(temp_file_path) as pdf:
                    metadata["page_count"] = len(pdf.pages)
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            # Extract text from page
                            page_text = page.extract_text()
                            
                            if page_text and page_text.strip():
                                extracted_text.append(page_text.strip())
                                metadata["successful_pages"] += 1
                            else:
                                # Try extracting text from tables if regular text extraction fails
                                tables = page.extract_tables()
                                if tables:
                                    metadata["tables_found"] += len(tables)
                                    table_text = []
                                    for table in tables:
                                        for row in table:
                                            if row:
                                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                                if row_text.strip():
                                                    table_text.append(row_text.strip())
                                    
                                    if table_text:
                                        extracted_text.append("\n".join(table_text))
                                        metadata["successful_pages"] += 1
                                    else:
                                        metadata["failed_pages"] += 1
                                else:
                                    logger.warning(f"No text or tables found on page {page_num} of {filename}")
                                    metadata["failed_pages"] += 1
                                    
                        except Exception as e:
                            logger.warning(f"Failed to extract from page {page_num} of {filename}: {str(e)}")
                            metadata["failed_pages"] += 1
                
                # Add PDF metadata if available
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    metadata.update({
                        "title": pdf.metadata.get("Title", ""),
                        "author": pdf.metadata.get("Author", ""),
                        "subject": pdf.metadata.get("Subject", ""),
                        "creator": pdf.metadata.get("Creator", ""),
                        "producer": pdf.metadata.get("Producer", "")
                    })
                
                if not extracted_text:
                    raise DocumentParsingError(f"No readable text found in PDF {filename} using pdfplumber")
                
                # Join all page text with double newlines
                full_text = "\n\n".join(extracted_text)
                
                # Clean up the text
                full_text = self._clean_extracted_text(full_text)
                
                return full_text, metadata
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
                    
        except Exception as e:
            logger.error(f"pdfplumber extraction failed for {filename}: {str(e)}")
            raise DocumentParsingError(f"pdfplumber extraction failed: {str(e)}")
    
    async def _extract_pdf_with_pypdf(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using pypdf (lightweight, good for simple PDFs)
        
        Args:
            content: PDF file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Create a BytesIO object from the content
            pdf_stream = io.BytesIO(content)
            
            # Create PDF reader
            pdf_reader = PdfReader(pdf_stream)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                # Try to decrypt with empty password (common case)
                try:
                    pdf_reader.decrypt("")
                except Exception:
                    raise DocumentParsingError(
                        f"PDF file {filename} is password-protected and cannot be processed"
                    )
            
            # Extract metadata
            metadata = {}
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                    "producer": pdf_reader.metadata.get("/Producer", ""),
                    "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")),
                    "modification_date": str(pdf_reader.metadata.get("/ModDate", ""))
                })
            
            metadata["page_count"] = len(pdf_reader.pages)
            
            # Extract text from all pages
            extracted_text = []
            successful_pages = 0
            failed_pages = 0
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        extracted_text.append(page_text.strip())
                        successful_pages += 1
                    else:
                        logger.warning(f"No text found on page {page_num} of {filename}")
                        failed_pages += 1
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num} of {filename}: {str(e)}")
                    failed_pages += 1
            
            metadata.update({
                "successful_pages": successful_pages,
                "failed_pages": failed_pages,
                "extraction_method": "pypdf"
            })
            
            if not extracted_text:
                raise DocumentParsingError(
                    f"No readable text found in PDF {filename} using pypdf"
                )
            
            # Join all page text with double newlines
            full_text = "\n\n".join(extracted_text)
            
            # Clean up the text
            full_text = self._clean_extracted_text(full_text)
            
            return full_text, metadata
            
        except DocumentParsingError:
            raise
        except PdfReadError as e:
            raise DocumentParsingError(f"Invalid or corrupted PDF file {filename}: {str(e)}")
        except Exception as e:
            logger.error(f"pypdf extraction failed for {filename}: {str(e)}")
            raise DocumentParsingError(f"pypdf extraction failed: {str(e)}")
    
    async def _extract_text_file(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from plain text file
        
        Args:
            content: Text file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Try UTF-8 first (most common)
            try:
                text = content.decode('utf-8')
                encoding = 'utf-8'
            except UnicodeDecodeError:
                # Try other common encodings
                encodings = ['latin-1', 'cp1252', 'iso-8859-1']
                text = None
                encoding = None
                
                for enc in encodings:
                    try:
                        text = content.decode(enc)
                        encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue
                
                if text is None:
                    raise DocumentParsingError(
                        f"Could not decode text file {filename}. "
                        "File may be corrupted or use an unsupported encoding."
                    )
            
            # Clean the text
            text = self._clean_extracted_text(text)
            
            metadata = {
                "encoding": encoding,
                "line_count": len(text.split('\n')),
                "extraction_method": "direct_decode"
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise DocumentParsingError(f"Failed to extract text from file: {str(e)}")
    
    async def _extract_docx_text(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file with comprehensive content extraction
        
        Args:
            content: DOCX file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        if not DOCX_AVAILABLE:
            raise DocumentParsingError(
                "DOCX processing not available. Please install python-docx library."
            )
        
        try:
            # Create a temporary file to work with python-docx
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Open document
                doc = Document(temp_file_path)
                
                # Initialize metadata
                metadata = {
                    "paragraph_count": 0,
                    "table_count": 0,
                    "header_count": 0,
                    "footer_count": 0,
                    "image_count": 0,
                    "hyperlink_count": 0,
                    "extraction_method": "python-docx",
                    "document_properties": {}
                }
                
                # Extract document properties
                if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                    metadata["document_properties"]["title"] = doc.core_properties.title
                if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                    metadata["document_properties"]["author"] = doc.core_properties.author
                if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                    metadata["document_properties"]["subject"] = doc.core_properties.subject
                if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                    metadata["document_properties"]["created"] = str(doc.core_properties.created)
                if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                    metadata["document_properties"]["modified"] = str(doc.core_properties.modified)
                
                # Extract text from paragraphs
                paragraphs = []
                hyperlink_count = 0
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                        
                        # Count hyperlinks in paragraph
                        for run in paragraph.runs:
                            if hasattr(run, '_element') and run._element.xpath('.//w:hyperlink'):
                                hyperlink_count += len(run._element.xpath('.//w:hyperlink'))
                
                metadata["paragraph_count"] = len(paragraphs)
                metadata["hyperlink_count"] = hyperlink_count
                
                # Extract text from tables with improved formatting
                table_text = []
                for table_idx, table in enumerate(doc.tables):
                    table_content = []
                    for row_idx, row in enumerate(table.rows):
                        row_text = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                # Handle multi-paragraph cells
                                cell_paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                                if len(cell_paragraphs) > 1:
                                    cell_text = " ".join(cell_paragraphs)
                                row_text.append(cell_text)
                            else:
                                row_text.append("")
                        
                        if any(cell.strip() for cell in row_text):  # Only add non-empty rows
                            table_content.append(" | ".join(row_text))
                    
                    if table_content:
                        # Add table header
                        table_text.append(f"[Table {table_idx + 1}]")
                        table_text.extend(table_content)
                        table_text.append("")  # Add spacing after table
                
                metadata["table_count"] = len(doc.tables)
                
                # Extract text from headers and footers
                header_footer_text = []
                
                # Headers
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                header_footer_text.append(f"[Header] {paragraph.text.strip()}")
                                metadata["header_count"] += 1
                    
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                header_footer_text.append(f"[Footer] {paragraph.text.strip()}")
                                metadata["footer_count"] += 1
                
                # Count images/shapes (approximate)
                try:
                    from docx.oxml.ns import qn
                    image_count = 0
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//a:blip'):
                                image_count += len(run._element.xpath('.//a:blip'))
                    metadata["image_count"] = image_count
                except Exception:
                    # If image counting fails, just continue
                    metadata["image_count"] = 0
                
                # Combine all text in logical order
                all_text_parts = []
                
                # Add headers first
                if header_footer_text:
                    headers = [text for text in header_footer_text if text.startswith("[Header]")]
                    if headers:
                        all_text_parts.extend(headers)
                        all_text_parts.append("")  # Add spacing
                
                # Add main content (paragraphs and tables interspersed)
                # This is a simplified approach - ideally we'd preserve the exact order
                if paragraphs:
                    all_text_parts.extend(paragraphs)
                
                if table_text:
                    if paragraphs:
                        all_text_parts.append("")  # Add spacing before tables
                    all_text_parts.extend(table_text)
                
                # Add footers last
                if header_footer_text:
                    footers = [text for text in header_footer_text if text.startswith("[Footer]")]
                    if footers:
                        all_text_parts.append("")  # Add spacing
                        all_text_parts.extend(footers)
                
                # Join all text
                full_text = "\n".join(all_text_parts)
                
                # Clean the text
                full_text = self._clean_extracted_text(full_text)
                
                # Validate that we extracted meaningful content
                if not full_text or len(full_text.strip()) < 10:
                    raise DocumentParsingError(
                        f"Insufficient readable content found in DOCX file {filename}. "
                        "The document may be empty, corrupted, or contain only images/objects."
                    )
                
                logger.info(f"Successfully extracted {len(full_text)} characters from DOCX {filename}")
                return full_text, metadata
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
            
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {filename}: {str(e)}")
            raise DocumentParsingError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_doc_text(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOC file (legacy Word format)
        
        This method attempts multiple approaches to extract text from DOC files:
        1. Check if it's actually a DOCX file with wrong extension
        2. Try basic text extraction for simple DOC files
        3. Provide helpful error message with conversion suggestions
        
        Args:
            content: DOC file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # First, check if this is actually a DOCX file with .doc extension
            # DOCX files start with PK (ZIP signature)
            if content.startswith(b'PK'):
                logger.info(f"File {filename} appears to be DOCX format despite .doc extension")
                try:
                    return await self._extract_docx_text(content, filename)
                except Exception as e:
                    logger.warning(f"Failed to process {filename} as DOCX: {str(e)}")
            
            # Check for DOC file signature
            if content.startswith(b'\xd0\xcf\x11\xe0'):
                logger.info(f"Detected legacy DOC format for {filename}")
                
                # Try basic text extraction for simple DOC files
                # This is a very basic approach that may work for simple documents
                try:
                    # Look for readable text in the binary content
                    # DOC files contain text mixed with binary data
                    text_content = self._extract_text_from_binary_doc(content)
                    
                    if text_content and len(text_content.strip()) >= 10:
                        metadata = {
                            "extraction_method": "binary_text_extraction",
                            "format": "legacy_doc",
                            "warning": "Basic extraction - some formatting and content may be missing"
                        }
                        
                        logger.info(f"Successfully extracted {len(text_content)} characters from DOC {filename} using basic method")
                        return text_content, metadata
                    
                except Exception as e:
                    logger.warning(f"Basic DOC extraction failed for {filename}: {str(e)}")
            
            # If all extraction methods fail, provide helpful error message
            raise DocumentParsingError(
                f"Legacy DOC format ({filename}) requires specialized processing. "
                f"For best results, please:\n"
                f"1. Convert to DOCX format using Microsoft Word or LibreOffice\n"
                f"2. Save as PDF for reliable text extraction\n"
                f"3. Copy and paste content into a text file\n\n"
                f"Alternative: Install additional tools like 'textract' or 'antiword' for DOC support."
            )
            
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error processing DOC file {filename}: {str(e)}")
            raise DocumentParsingError(
                f"Failed to process DOC file {filename}. "
                f"Please convert to DOCX or PDF format for reliable text extraction."
            )
    
    def _extract_text_from_binary_doc(self, content: bytes) -> str:
        """
        Attempt basic text extraction from binary DOC content
        
        This is a simple heuristic approach that looks for readable text
        in the binary DOC file. It's not comprehensive but may work for
        simple documents.
        
        Args:
            content: Binary DOC file content
            
        Returns:
            Extracted text string
        """
        try:
            # Convert bytes to string, ignoring non-text characters
            # This is a very basic approach
            text_parts = []
            
            # Try to decode as various encodings and extract readable text
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    
                    # Extract sequences of printable characters
                    import re
                    # Look for sequences of at least 3 printable characters
                    text_sequences = re.findall(r'[a-zA-Z0-9\s\.,;:!?\-\'"()]{3,}', decoded)
                    
                    # Filter out sequences that are likely binary data
                    meaningful_sequences = []
                    for seq in text_sequences:
                        # Skip sequences with too many special characters or numbers
                        if len(seq.strip()) >= 3:
                            alpha_ratio = sum(c.isalpha() or c.isspace() for c in seq) / len(seq)
                            if alpha_ratio > 0.5:  # At least 50% letters or spaces
                                meaningful_sequences.append(seq.strip())
                    
                    if meaningful_sequences:
                        text_parts.extend(meaningful_sequences)
                        break  # Use the first encoding that gives good results
                        
                except UnicodeDecodeError:
                    continue
            
            if text_parts:
                # Join and clean the extracted text
                full_text = " ".join(text_parts)
                full_text = self._clean_extracted_text(full_text)
                
                # Remove excessive repetition (common in binary extraction)
                lines = full_text.split('\n')
                unique_lines = []
                seen_lines = set()
                
                for line in lines:
                    line = line.strip()
                    if line and line not in seen_lines and len(line) > 2:
                        unique_lines.append(line)
                        seen_lines.add(line)
                
                return '\n'.join(unique_lines)
            
            return ""
            
        except Exception as e:
            logger.warning(f"Binary DOC text extraction failed: {str(e)}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace while preserving paragraph structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Strip whitespace from each line
            cleaned_line = line.strip()
            cleaned_lines.append(cleaned_line)
        
        # Join lines back together
        text = '\n'.join(cleaned_lines)
        
        # Remove excessive blank lines (more than 2 consecutive)
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_supported_formats(self) -> list:
        """
        Get list of supported file formats
        
        Returns:
            List of supported file extensions
        """
        return list(self.supported_formats.keys())
    
    def is_format_supported(self, file_extension: str) -> bool:
        """
        Check if a file format is supported
        
        Args:
            file_extension: File extension (e.g., '.pdf')
            
        Returns:
            True if format is supported, False otherwise
        """
        return file_extension.lower() in self.supported_formats